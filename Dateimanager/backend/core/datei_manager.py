import os
import sys # sys hinzufügen für Plattformprüfung
import json
import csv
import subprocess
import logging
import time
from docx import Document
from PyPDF2 import PdfWriter # Beachte: PyPDF2 kann keine PDF *aus Text* erstellen, nur leere oder aus anderen PDFs
from typing import Optional, Dict, List
import platform # platform hinzufügen für OS-spezifische Logik

# --- Hilfsfunktionen für plattformunabhängiges Öffnen ---

def open_path_cross_platform(path: str):
    """Öffnet einen Pfad (Datei oder Ordner) mit dem Standardprogramm des Systems."""
    try:
        if platform.system() == "Windows":
            # os.startfile ist oft zuverlässiger als 'start'
            os.startfile(os.path.normpath(path))
        elif platform.system() == "Darwin": # macOS
            subprocess.run(['open', path], check=True)
        else: # Linux und andere Unix-ähnliche
            subprocess.run(['xdg-open', path], check=True)
        return True, f"'{path}' erfolgreich zum Öffnen übergeben."
    except FileNotFoundError:
         # Fallback für Linux, falls xdg-open nicht gefunden wird
         if platform.system() != "Windows" and platform.system() != "Darwin":
             try:
                 subprocess.run(['gvfs-open', path], check=True) # Für GNOME
                 return True, f"'{path}' erfolgreich zum Öffnen übergeben (via gvfs-open)."
             except FileNotFoundError:
                  try:
                      subprocess.run(['kde-open', path], check=True) # Für KDE
                      return True, f"'{path}' erfolgreich zum Öffnen übergeben (via kde-open)."
                  except FileNotFoundError:
                      logging.error(f"Konnte kein Standardprogramm zum Öffnen von '{path}' finden (xdg-open, gvfs-open, kde-open nicht gefunden).")
                      return False, f"Kein Standardprogramm zum Öffnen von '{path}' gefunden."
             except Exception as e:
                 logging.error(f"Fehler beim Öffnen von '{path}' mit Fallback: {e}")
                 return False, f"Fehler beim Öffnen von '{path}': {e}"
         else:
            logging.error(f"Fehler: Pfad '{path}' nicht gefunden.")
            return False, f"Fehler: Pfad '{path}' nicht gefunden."
    except Exception as e:
        logging.error(f"Fehler beim Öffnen von '{path}': {e}")
        return False, f"Fehler beim Öffnen von '{path}': {e}"

def reveal_in_file_manager(path: str):
    """Öffnet den Dateimanager und zeigt den angegebenen Pfad an (oder dessen Ordner)."""
    try:
        norm_path = os.path.normpath(path)
        if not os.path.exists(norm_path):
             return False, f"Fehler: Pfad '{norm_path}' existiert nicht."

        if platform.system() == "Windows":
            # Öffnet den Explorer und wählt die Datei/den Ordner aus
            subprocess.run(['explorer', '/select,', norm_path], check=True)
        elif platform.system() == "Darwin": # macOS
            # Öffnet den Finder und zeigt die Datei/den Ordner an
            subprocess.run(['open', '-R', norm_path], check=True)
        else: # Linux
            # Öffnet den Dateimanager im übergeordneten Ordner
            # Auswahl der Datei ist nicht standardisiert
            dir_path = os.path.dirname(norm_path) if os.path.isfile(norm_path) else norm_path
            subprocess.run(['xdg-open', dir_path], check=True)
        return True, f"'{norm_path}' im Dateimanager angezeigt."
    except Exception as e:
        logging.error(f"Fehler beim Anzeigen von '{path}' im Dateimanager: {e}")
        return False, f"Fehler beim Anzeigen von '{path}' im Dateimanager: {e}"


# --- DateiManager Klasse ---

class DateiManager:
    def __init__(self, structure_file="data/structure.json"):
        # Systemordner - erweitert für Cross-Platform (Beispiele)
        self.system_folders = self._get_platform_system_folders()
        self.structure_file = structure_file
        os.makedirs(os.path.dirname(self.structure_file), exist_ok=True)
        self.cached_structure = self._load_structure()

    def _get_platform_system_folders(self):
        """Gibt eine Liste von typischen System-/Ignorier-Ordnern für das aktuelle OS zurück."""
        folders = {
            "$Recycle.Bin", "$RECYCLE.BIN", "System Volume Information", "Recovery", # Gemeinsam?
            "node_modules", ".git", ".vscode", "__pycache__", ".venv", "venv", # Entwicklung
            ".DS_Store", ".localized" # macOS
        }
        system = platform.system()
        if system == "Windows":
            folders.update([
                "Windows", "ProgramData", "Program Files", "Program Files (x86)",
                "AppData", "Users", # Users kann je nach Anwendungsfall drin bleiben
                "Dokumente und Einstellungen" # Älter
            ])
        elif system == "Darwin": # macOS
            folders.update([
                "Applications", "Library", "System", "User Information", "Volumes",
                "private", "cores", "opt", "usr", "bin", "sbin", "etc", "var" # Typische Unix-Ordner im Root
            ])
        else: # Linux
            folders.update([
                "proc", "sys", "dev", "run", "lost+found",
                "snap", # Ubuntu Snap-Ordner
                "opt", "usr", "bin", "sbin", "etc", "var", "lib", "lib64", "boot", "root" # Typische Root-Ordner
            ])
        return folders

    def _load_structure(self):
        # (Keine Änderung nötig)
        try:
            with open(self.structure_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            logging.info(f"Strukturdatei '{self.structure_file}' nicht gefunden. Wird beim ersten Scan erstellt.")
            return None
        except json.JSONDecodeError:
            logging.warning(f"Fehler beim Dekodieren der Strukturdatei. Starte mit leerer Struktur.")
            return None
        except Exception as e:
            logging.error(f"Fehler beim Laden der Struktur: {e}")
            return None

    def _save_structure(self, structure):
        # (Keine Änderung nötig)
        try:
            with open(self.structure_file, "w", encoding="utf-8") as f:
                json.dump(structure, f, indent=4)
            logging.info(f"Struktur erfolgreich in '{self.structure_file}' gespeichert.")
            return True
        except Exception as e:
            logging.error(f"Fehler beim Speichern der Struktur: {e}")
            return False

    def open_file_or_folder(self, path: str) -> str:
        """Öffnet eine Datei oder einen Ordner mit dem Systemstandard (plattformunabhängig)."""
        success, message = open_path_cross_platform(path)
        return message # Gibt die Erfolgs-/Fehlermeldung zurück

    def open_in_explorer(self, file_path: str) -> str:
        """Zeigt die Datei oder den Ordner im nativen Dateimanager an (plattformunabhängig)."""
        success, message = reveal_in_file_manager(file_path)
        # bring_explorer_to_foreground entfernt, da nicht portabel
        return message

    # --- Schreibfunktionen ---
    # (Keine direkten Änderungen nötig, da os.path plattformunabhängig ist)
    # ABER: _write_pdf ist problematisch, siehe Kommentar unten.

    def write_content(self, file_path: str, content: str) -> bool:
        ext = os.path.splitext(file_path)[-1].lower()
        try:
            if ext in {".txt", ".md", ".html", ".css", ".js", ".py", ".xml"}:
                return self._write_text_file(file_path, content)
            elif ext == ".json":
                return self._write_json(file_path, content)
            elif ext == ".csv":
                return self._write_csv(file_path, content)
            elif ext == ".docx":
                return self._write_docx(file_path, content)
            elif ext == ".pdf":
                # ACHTUNG: Diese Funktion erstellt nur eine LEERE PDF-Seite.
                # Sie schreibt NICHT den 'content' in die PDF.
                # PDF-Erstellung aus Text ist komplexer (benötigt z.B. reportlab, fpdf2).
                logging.warning(f"Hinweis: _write_pdf erstellt nur eine leere PDF für '{file_path}', schreibt keinen Inhalt.")
                return self._write_pdf_placeholder(file_path) # Umbenannt zur Klarheit
        except Exception as e:
            logging.error(f"Fehler beim Schreiben von '{file_path}': {e}")
            return False
        logging.warning(f"Dateityp '{ext}' für '{file_path}' wird nicht unterstützt zum Schreiben.")
        return False

    def _write_text_file(self, path, content):
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            logging.info(f"Textdatei '{path}' erfolgreich geschrieben.")
            return True
        except Exception as e:
            logging.error(f"Fehler beim Schreiben der Textdatei '{path}': {e}")
            return False

    def _write_json(self, path, content):
        try:
            # Versuche, den String als JSON zu parsen, um die Validität zu prüfen
            data = json.loads(content)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
            logging.info(f"JSON-Datei '{path}' erfolgreich geschrieben.")
            return True
        except json.JSONDecodeError as e:
            logging.error(f"Ungültiges JSON-Format beim Schreiben von '{path}': {e}")
            return False
        except Exception as e:
            logging.error(f"Fehler beim Schreiben der JSON-Datei '{path}': {e}")
            return False

    def _write_csv(self, path, content):
        try:
            # Annahme: content ist ein String, jede Zeile ist eine CSV-Zeile
            # Besser wäre es, eine Liste von Listen zu erwarten.
            # Diese Implementierung ist sehr einfach gehalten.
            lines = content.strip().split("\n")
            # Versuche, jede Zeile zu parsen (einfache Komma-Trennung)
            rows = [line.split(",") for line in lines]
            with open(path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            logging.info(f"CSV-Datei '{path}' erfolgreich geschrieben.")
            return True
        except Exception as e:
            logging.error(f"Fehler beim Schreiben der CSV-Datei '{path}': {e}")
            return False

    def _write_docx(self, path, content):
        try:
            doc = Document()
            # Füge den Inhalt als einen einzigen Absatz hinzu (vereinfacht)
            doc.add_paragraph(content)
            doc.save(path)
            logging.info(f"DOCX-Datei '{path}' erfolgreich geschrieben.")
            return True
        except Exception as e:
            logging.error(f"Fehler beim Schreiben der DOCX-Datei '{path}': {e}")
            return False

    def _write_pdf_placeholder(self, path):
        """Erstellt eine leere A4-PDF-Seite (nur Platzhalter)."""
        try:
            writer = PdfWriter()
            # Fügt eine leere A4-Seite hinzu (Dimensionen in Punkten: 595.27 x 841.89)
            writer.add_blank_page()
            with open(path, "wb") as f:
                writer.write(f)
            logging.info(f"Leere PDF-Platzhalterdatei '{path}' erstellt.")
            return True
        except Exception as e:
            logging.error(f"Fehler beim Erstellen der PDF-Platzhalterdatei '{path}': {e}")
            return False

    def delete_file(self, file_path: str) -> bool:
        # (Keine Änderung nötig)
        try:
            norm_path = os.path.normpath(file_path)
            if os.path.isfile(norm_path):
                 os.remove(norm_path)
                 logging.info(f"Datei '{norm_path}' erfolgreich gelöscht.")
                 return True
            else:
                 logging.warning(f"Löschen fehlgeschlagen: '{norm_path}' ist keine Datei oder existiert nicht.")
                 return False
        except Exception as e:
            logging.error(f"Fehler beim Löschen der Datei '{file_path}': {e}")
            return False

    def get_file_structure(self, path: str = None, force_rescan=False):
        """
        Gibt die Verzeichnisstruktur zurück, entweder aus dem Cache oder durch Neuscan.
        Passt sich an das Betriebssystem an (Laufwerke vs. Root-Verzeichnis).
        """
        if self.cached_structure is None or force_rescan:
            logging.info(f"Struktur-Cache {'leer' if self.cached_structure is None else 'wird neu erstellt'} (force_rescan={force_rescan}). Starte Scan...")
            if path is None:
                # Startpunkte basierend auf OS bestimmen
                if platform.system() == "Windows":
                    # Liste verfügbare Laufwerke auf
                    drives = [d + ":\\" for d in "ABCDEFGHIJKLMNOPQRSTUVWXYZ" if os.path.exists(d + ":\\")]
                    logging.info(f"Scanne Windows-Laufwerke: {drives}")
                    # Struktur für jedes Laufwerk holen
                    self.cached_structure = [{"name": drive, "path": drive, "children": self._get_directory_structure(drive)} for drive in drives]
                else:
                    # Für Linux/macOS: Starte im Wurzelverzeichnis '/'
                    root_path = "/"
                    logging.info(f"Scanne Unix-Wurzelverzeichnis: {root_path}")
                    # Die Struktur ist direkt die Liste der Kinder von '/'
                    # Wir geben eine Liste zurück, die das Root-Element enthält
                    self.cached_structure = [{"name": "/", "path": "/", "children": self._get_directory_structure(root_path)}]
            else:
                # Wenn ein spezifischer Pfad gegeben ist
                norm_path = os.path.normpath(path)
                logging.info(f"Scanne spezifischen Pfad: {norm_path}")
                if not os.path.exists(norm_path) or not os.path.isdir(norm_path):
                    logging.error(f"Angegebener Pfad '{norm_path}' existiert nicht oder ist kein Verzeichnis.")
                    return None # Oder leere Struktur? Je nach Anforderung.
                # Gib nur die Struktur für diesen Pfad zurück
                # Wichtig: Das Ergebnis von _get_directory_structure ist eine Liste von Kindern.
                # Für Konsistenz mit dem Laufwerk-Scan geben wir ein Objekt zurück.
                self.cached_structure = [{"name": os.path.basename(norm_path) or norm_path, # Name des Ordners oder Pfad selbst
                                          "path": norm_path,
                                          "children": self._get_directory_structure(norm_path)}]

            # Speichere die neu gescannte Struktur (wenn erfolgreich)
            if self.cached_structure is not None:
                 self._save_structure(self.cached_structure)
            else:
                 logging.error("Konnte keine gültige Struktur scannen.")

        # Gib die gecachte (oder gerade gescannte) Struktur zurück
        # Wenn ein spezifischer Pfad angefragt wurde, aber kein Cache existierte,
        # könnte self.cached_structure jetzt die Struktur *nur* für diesen Pfad enthalten.
        # Wenn der Aufrufer die *gesamte* Struktur erwartet, muss das angepasst werden.
        # Aktuell: Gibt immer den *gesamten* Cache zurück, auch wenn nur ein Teil gescannt wurde.
        return self.cached_structure

    def _get_directory_structure(self, path: str, depth=0, max_depth=5): # Tiefenbegrenzung hinzugefügt
        """Rekursive Hilfsfunktion zum Scannen der Verzeichnisstruktur."""
        # Verhindere zu tiefe Rekursion
        if depth > max_depth:
            logging.warning(f"Maximale Rekursionstiefe ({max_depth}) erreicht bei Pfad: {path}")
            return []

        dir_structure = []
        try:
            # Verwende try-except um os.scandir für Berechtigungsprobleme etc.
            with os.scandir(path) as entries:
                for entry in entries:
                    # Ignoriere Systemordner und versteckte Ordner (beginnen mit .)
                    is_hidden = entry.name.startswith('.')
                    is_system = entry.name in self.system_folders
                    if entry.is_dir(follow_symlinks=False) and not is_hidden and not is_system:
                        try:
                            # Rekursiver Aufruf für Unterordner
                            children = self._get_directory_structure(entry.path, depth + 1, max_depth)
                            dir_structure.append({
                                "name": entry.name,
                                "path": os.path.normpath(entry.path), # Normierten Pfad speichern
                                "children": children
                            })
                        except PermissionError:
                             logging.warning(f"Keine Berechtigung für Ordner: {entry.path}")
                        except Exception as e_inner:
                             logging.error(f"Fehler beim Verarbeiten des Unterordners '{entry.path}': {e_inner}")

        except PermissionError:
            logging.warning(f"Keine Berechtigung für Ordner: {path}")
        except FileNotFoundError:
             logging.warning(f"Ordner nicht gefunden während Scan (evtl. gelöscht?): {path}")
        except Exception as e:
            logging.error(f"Allgemeiner Fehler beim Auslesen der Ordnerstruktur für '{path}': {e}")

        return dir_structure

    def rescan_file_structure(self, path: str = None):
        """Erzwingt einen Neuscan der Ordnerstruktur und speichert das Ergebnis."""
        logging.info(f"Erzwinge Neuscan für {'alle Startpunkte' if path is None else path}...")
        # Setze Cache zurück und rufe get_file_structure mit force_rescan=True auf
        self.cached_structure = None
        return self.get_file_structure(path=path, force_rescan=True)

    # --- Lesefunktionen (unverändert, nutzen os.path) ---
    def read_file(self, file_path: str) -> Optional[str]:
        ext = os.path.splitext(file_path)[-1].lower()
        norm_path = os.path.normpath(file_path) # Normierten Pfad verwenden
        try:
            if not os.path.isfile(norm_path):
                 logging.warning(f"Datei nicht gefunden oder kein File: '{norm_path}'")
                 return None

            if ext in {".txt", ".md", ".html", ".css", ".js", ".py", ".xml"}:
                with open(norm_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif ext == ".json":
                with open(norm_path, "r", encoding="utf-8") as f:
                    # Gib den rohen String zurück, damit er im Frontend ggf. formatiert werden kann
                    return f.read()
                    # return json.dumps(json.load(f), indent=4) # Alternative: Formatiert zurückgeben
            elif ext == ".csv":
                with open(norm_path, "r", newline="", encoding="utf-8") as f:
                    # Gib CSV als String zurück, Zeilen mit \n getrennt
                    return f.read()
                    # reader = csv.reader(f)
                    # return "\n".join([",".join(row) for row in reader]) # Alternative
            elif ext == ".docx":
                 try:
                    doc = Document(norm_path)
                    return "\n".join([paragraph.text for paragraph in doc.paragraphs])
                 except Exception as docx_e:
                     logging.error(f"Fehler beim Lesen von DOCX '{norm_path}' mit python-docx: {docx_e}")
                     return f"Fehler beim Lesen von DOCX: {docx_e}" # Gib Fehlermeldung zurück
            elif ext == ".pdf":
                # Hier könnte man eine Text-Extraktionsbibliothek einbinden (z.B. pypdf oder pymupdf)
                # Fürs Erste geben wir eine Meldung zurück.
                return "Vorschau für PDF-Inhalt nicht direkt unterstützt. OCR verwenden?"
            else:
                logging.info(f"Dateityp '{ext}' für '{norm_path}' wird nicht für Inhaltsvorschau unterstützt.")
                return None # Oder eine Meldung wie "Vorschau nicht unterstützt"
        except Exception as e:
            logging.error(f"Fehler beim Lesen der Datei '{norm_path}': {e}")
            # Gib spezifischere Fehlermeldung zurück, wenn möglich
            return f"Fehler beim Lesen der Datei: {e}"


    def overwrite_file(self, file_path: str, content: str) -> bool:
        """Überschreibt den Inhalt einer Datei mit dem gegebenen Inhalt."""
        # Nutzt die aktualisierte write_content Funktion
        return self.write_content(file_path, content)

    # --- JSON/Text Lese/Schreib-Wrapper (unverändert) ---
    def read_json_file(self, file_path: str):
        norm_path = os.path.normpath(file_path)
        content_str = self.read_file(norm_path) # Nutzt die zentrale Lesefunktion
        if content_str and not content_str.startswith("Fehler"):
            try:
                return json.loads(content_str)
            except json.JSONDecodeError as e:
                logging.error(f"Fehler beim Dekodieren von JSON aus '{norm_path}': {e}")
                return {"error": f"Fehler beim Dekodieren der JSON-Datei: {e}"}
        elif content_str and content_str.startswith("Fehler"):
             return {"error": content_str} # Gib Lesefehler weiter
        else: # content is None oder leer
             return {"error": f"Datei '{norm_path}' konnte nicht gelesen werden oder ist leer."}


    def save_json_file(self, file_path: str, data: Dict):
        try:
            # Konvertiere Python-Dict zu JSON-String
            content = json.dumps(data, indent=4, ensure_ascii=False) # ensure_ascii=False für Umlaute etc.
            if self.write_content(file_path, content): # Nutzt zentrale Schreibfunktion
                return {"message": f"Datei '{file_path}' erfolgreich gespeichert."}
            else:
                # write_content loggt den Fehler bereits
                return {"error": f"Fehler beim Speichern der JSON-Datei '{file_path}'."}
        except TypeError as e:
             logging.error(f"Fehler beim Konvertieren der Daten zu JSON für '{file_path}': {e}")
             return {"error": f"Daten konnten nicht als JSON gespeichert werden: {e}"}
        except Exception as e:
             logging.error(f"Allgemeiner Fehler beim Speichern von JSON '{file_path}': {e}")
             return {"error": f"Allgemeiner Fehler beim Speichern der JSON-Datei '{file_path}'."}


    def read_text_file(self, file_path: str):
        content = self.read_file(file_path) # Nutzt zentrale Lesefunktion
        if content and not content.startswith("Fehler"):
            return {"content": content}
        elif content and content.startswith("Fehler"):
            return {"error": content} # Gib Lesefehler weiter
        else:
            return {"error": f"Datei '{file_path}' konnte nicht gelesen werden oder ist leer."}

    def save_text_file(self, file_path: str, content: str):
        if self.write_content(file_path, content): # Nutzt zentrale Schreibfunktion
            return {"message": f"Datei '{file_path}' erfolgreich gespeichert."}
        else:
            # write_content loggt den Fehler bereits
            return {"error": f"Fehler beim Speichern der Textdatei '{file_path}'."}